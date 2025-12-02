"""
Report Generator Module

Generates consolidated reports in DOCX and PDF formats.
Includes all scan results, team information, and visualizations.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  WARNING: python-docx not available. Install with: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  WARNING: reportlab not available. Install with: pip install reportlab")

from logger import Logger
from config import Config
from identity_checker import IdentityChecker


class ReportGenerator:
    """Generates consolidated security assessment reports."""
    
    def __init__(self, logger: Logger, config: Config, identity_checker: IdentityChecker):
        """
        Initialize the report generator.
        
        Args:
            logger: Logger instance
            config: Configuration instance
            identity_checker: Identity checker instance
        """
        self.logger = logger
        self.config = config
        self.identity_checker = identity_checker
    
    def generate_report(self, input_dir: str, output_file: Optional[str] = None,
                       format_type: str = 'both') -> int:
        """
        Generate consolidated report from scan results.
        
        Args:
            input_dir: Directory containing scan result files
            output_file: Optional output file path
            format_type: Report format ('docx', 'pdf', or 'both')
            
        Returns:
            Exit code (0 for success)
        """
        input_path = Path(input_dir)
        if not input_path.exists():
            print(f"❌ ERROR: Input directory not found: {input_dir}")
            self.logger.error(f"Input directory not found: {input_dir}")
            return 1
        
        self.logger.info(f"Generating report from {input_dir}")
        self.logger.log_operation('report_generator', 'generate_start', {
            'input_dir': input_dir,
            'format': format_type
        })
        
        # Collect all scan results
        results = self._collect_results(input_path)
        
        if not results:
            print("⚠️  WARNING: No scan results found in input directory")
            self.logger.warning("No scan results found")
            return 1
        
        # Get team information
        try:
            team_info = self.identity_checker.get_team_info()
        except Exception:
            team_info = {'team_name': 'Unknown', 'members': []}
        
        # Generate output path
        if output_file:
            output_path = Path(output_file)
        else:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.config.get_output_dir() / f"report_{timestamp}"
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Generate reports
        if format_type in ['docx', 'both']:
            if DOCX_AVAILABLE:
                docx_file = output_path.with_suffix('.docx')
                self._generate_docx(docx_file, team_info, results)
                print(f"✅ DOCX report generated: {docx_file}")
            else:
                print("⚠️  WARNING: python-docx not available, skipping DOCX generation")
        
        if format_type in ['pdf', 'both']:
            if PDF_AVAILABLE:
                pdf_file = output_path.with_suffix('.pdf')
                self._generate_pdf(pdf_file, team_info, results)
                print(f"✅ PDF report generated: {pdf_file}")
            else:
                print("⚠️  WARNING: reportlab not available, skipping PDF generation")
        
        # Generate JSON summary
        json_file = output_path.with_suffix('.json')
        self._generate_json_summary(json_file, team_info, results)
        print(f"✅ JSON summary generated: {json_file}")
        
        return 0
    
    def _collect_results(self, input_path: Path) -> Dict:
        """
        Collect all scan results from input directory.
        
        Args:
            input_path: Input directory path
            
        Returns:
            Dictionary of collected results
        """
        results = {
            'port_scans': [],
            'password_tests': [],
            'stress_tests': [],
            'web_discoveries': [],
            'packet_captures': []
        }
        
        # Find all JSON files
        for json_file in input_path.glob('*.json'):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    
                    # Categorize by filename pattern
                    filename = json_file.name.lower()
                    if 'scan_' in filename:
                        results['port_scans'].append(data)
                    elif 'auth_test' in filename:
                        results['password_tests'].append(data)
                    elif 'stress_' in filename:
                        results['stress_tests'].append(data)
                    elif 'footprint_' in filename:
                        results['web_discoveries'].append(data)
                    elif 'pcap_' in filename:
                        results['packet_captures'].append(data)
                        
            except Exception as e:
                self.logger.debug(f"Error reading {json_file}: {e}")
        
        return results
    
    def _generate_docx(self, output_file: Path, team_info: Dict, results: Dict) -> None:
        """
        Generate DOCX report.
        
        Args:
            output_file: Output file path
            team_info: Team information dictionary
            results: Collected scan results
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Finsecure Security Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('PayBuddy Security Assessment')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.italic = True
        
        doc.add_paragraph()
        
        # Warning
        warning = doc.add_paragraph(
            '⚠️  AUTHORIZED EDUCATIONAL USE ONLY - For controlled lab environments only',
            style='Intense Quote'
        )
        if warning.runs:
            warning_format = warning.runs[0].font
            warning_format.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_page_break()
        
        # Team Information
        doc.add_heading('Team Information', 1)
        doc.add_paragraph(f"Team Name: {team_info.get('team_name', 'N/A')}")
        doc.add_paragraph(f"Members: {', '.join(team_info.get('members', []))}")
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_page_break()
        
        # Port Scan Results
        if results['port_scans']:
            doc.add_heading('Port Scan Results', 1)
            for scan in results['port_scans']:
                doc.add_heading(f"Target: {scan.get('target', 'N/A')}", 2)
                doc.add_paragraph(f"Scan Time: {scan.get('timestamp', 'N/A')}")
                doc.add_paragraph(f"Open Ports Found: {scan.get('open_count', 0)}")
                
                if scan.get('open_ports'):
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Port'
                    hdr_cells[1].text = 'Service'
                    hdr_cells[2].text = 'Banner'
                    
                    for port_info in scan['open_ports']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(port_info.get('port', 'N/A'))
                        row_cells[1].text = port_info.get('service', 'N/A')
                        row_cells[2].text = str(port_info.get('banner', 'N/A'))[:50]
                
                doc.add_paragraph()
        
        # Password Test Results
        if results['password_tests']:
            doc.add_heading('Password Assessment Results', 1)
            for test in results['password_tests']:
                doc.add_paragraph(f"Test Time: {test.get('timestamp', 'N/A')}")
                doc.add_paragraph(f"Total Passwords: {test.get('total_passwords', 0)}")
                doc.add_paragraph(f"Strong: {test.get('strong_count', 0)}")
                doc.add_paragraph(f"Moderate: {test.get('moderate_count', 0)}")
                doc.add_paragraph(f"Weak: {test.get('weak_count', 0)}")
                doc.add_paragraph(f"Average Entropy: {test.get('average_entropy', 0):.2f} bits")
                doc.add_paragraph()
        
        # Stress Test Results
        if results['stress_tests']:
            doc.add_heading('Load Test Results', 1)
            for test in results['stress_tests']:
                doc.add_heading(f"Target: {test.get('target', 'N/A')}", 2)
                doc.add_paragraph(f"Test Time: {test.get('timestamp', 'N/A')}")
                doc.add_paragraph(f"Total Requests: {test.get('total_requests', 0)}")
                doc.add_paragraph(f"Successful: {test.get('successful_requests', 0)}")
                doc.add_paragraph(f"Failed: {test.get('failed_requests', 0)}")
                doc.add_paragraph(f"Average Response Time: {test.get('avg_response_time', 0):.2f}ms")
                doc.add_paragraph(f"Requests/Second: {test.get('requests_per_second', 0):.2f}")
                doc.add_paragraph()
        
        # Web Discovery Results
        if results['web_discoveries']:
            doc.add_heading('Web Discovery Results', 1)
            for discovery in results['web_discoveries']:
                doc.add_heading(f"Target: {discovery.get('target', 'N/A')}", 2)
                doc.add_paragraph(f"Directories Found: {discovery.get('directories_found', 0)}")
                doc.add_paragraph(f"Subdomains Found: {discovery.get('subdomains_found', 0)}")
                
                if discovery.get('directories'):
                    doc.add_paragraph("Found Directories:")
                    for dir_info in discovery['directories'][:20]:  # First 20
                        doc.add_paragraph(f"  • {dir_info.get('url', 'N/A')} ({dir_info.get('status_code', 'N/A')})",
                                         style='List Bullet')
                
                if discovery.get('subdomains'):
                    doc.add_paragraph("Found Subdomains:")
                    for subdomain_info in discovery['subdomains']:
                        doc.add_paragraph(f"  • {subdomain_info.get('subdomain', 'N/A')}",
                                         style='List Bullet')
                
                doc.add_paragraph()
        
        # Packet Capture Results
        if results['packet_captures']:
            doc.add_heading('Packet Capture Results', 1)
            for capture in results['packet_captures']:
                doc.add_paragraph(f"Capture Time: {capture.get('timestamp', 'N/A')}")
                doc.add_paragraph(f"Total Packets: {capture.get('total_packets', 0)}")
                doc.add_paragraph(f"TCP: {capture.get('tcp_count', 0)}")
                doc.add_paragraph(f"UDP: {capture.get('udp_count', 0)}")
                doc.add_paragraph(f"ICMP: {capture.get('icmp_count', 0)}")
                doc.add_paragraph()
        
        # Footer
        doc.add_page_break()
        footer = doc.add_paragraph('Generated by Finsecure Toolkit - Educational Use Only')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_format = footer.runs[0].font
        footer_format.size = Pt(10)
        footer_format.italic = True
        
        doc.save(str(output_file))
        self.logger.info(f"DOCX report saved to {output_file}")
    
    def _generate_pdf(self, output_file: Path, team_info: Dict, results: Dict) -> None:
        """
        Generate PDF report.
        
        Args:
            output_file: Output file path
            team_info: Team information dictionary
            results: Collected scan results
        """
        doc = SimpleDocTemplate(str(output_file), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph('<b>Finsecure Security Assessment Report</b>', styles['Title'])
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        subtitle = Paragraph('<i>PayBuddy Security Assessment</i>', styles['Normal'])
        story.append(subtitle)
        story.append(Spacer(1, 0.2*inch))
        
        # Warning
        warning = Paragraph(
            '<font color="red"><b>⚠️  AUTHORIZED EDUCATIONAL USE ONLY</b></font><br/>'
            'For controlled lab environments only',
            styles['Normal']
        )
        story.append(warning)
        story.append(PageBreak())
        
        # Team Information
        story.append(Paragraph('<b>Team Information</b>', styles['Heading1']))
        story.append(Paragraph(f"Team Name: {team_info.get('team_name', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"Members: {', '.join(team_info.get('members', []))}", styles['Normal']))
        story.append(Paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(PageBreak())
        
        # Port Scan Results
        if results['port_scans']:
            story.append(Paragraph('<b>Port Scan Results</b>', styles['Heading1']))
            for scan in results['port_scans']:
                story.append(Paragraph(f"<b>Target:</b> {scan.get('target', 'N/A')}", styles['Heading2']))
                story.append(Paragraph(f"Open Ports Found: {scan.get('open_count', 0)}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Add more sections as needed...
        
        # Footer
        story.append(PageBreak())
        footer = Paragraph('<i>Generated by Finsecure Toolkit - Educational Use Only</i>', styles['Normal'])
        story.append(footer)
        
        doc.build(story)
        self.logger.info(f"PDF report saved to {output_file}")
    
    def _generate_json_summary(self, output_file: Path, team_info: Dict, results: Dict) -> None:
        """
        Generate JSON summary report.
        
        Args:
            output_file: Output file path
            team_info: Team information dictionary
            results: Collected scan results
        """
        summary = {
            'report_metadata': {
                'generated_at': datetime.now().isoformat(),
                'team_name': team_info.get('team_name', 'N/A'),
                'members': team_info.get('members', [])
            },
            'summary': {
                'port_scans': len(results['port_scans']),
                'password_tests': len(results['password_tests']),
                'stress_tests': len(results['stress_tests']),
                'web_discoveries': len(results['web_discoveries']),
                'packet_captures': len(results['packet_captures'])
            },
            'results': results
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)
        
        self.logger.info(f"JSON summary saved to {output_file}")

